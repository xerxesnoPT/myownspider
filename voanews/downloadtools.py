import pymongo
from io import BytesIO
from PIL import Image
from docx.shared import Inches
from docx import Document
import requests
import re
import random

class Writeto_docx(object):
    def __init__(self):
        self.mongo_collection = pymongo.MongoClient()['news']['info']
        self.ualist = [
            "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/22.0.1207.1 Safari/537.1",
            "Mozilla/5.0 (X11; CrOS i686 2268.111.0) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.57 Safari/536.11",
            "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1092.0 Safari/536.6",
            "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.6 (KHTML, like Gecko) Chrome/20.0.1090.0 Safari/536.6",
            "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/19.77.34.5 Safari/537.1",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.9 Safari/536.5",
            "Mozilla/5.0 (Windows NT 6.0) AppleWebKit/536.5 (KHTML, like Gecko) Chrome/19.0.1084.36 Safari/536.5",
            "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
            "Mozilla/5.0 (Windows NT 5.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_8_0) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1063.0 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1062.0 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.1) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.1 Safari/536.3",
            "Mozilla/5.0 (Windows NT 6.2) AppleWebKit/536.3 (KHTML, like Gecko) Chrome/19.0.1061.0 Safari/536.3",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24",
            "Mozilla/5.0 (Windows NT 6.2; WOW64) AppleWebKit/535.24 (KHTML, like Gecko) Chrome/19.0.1055.1 Safari/535.24"
        ]

    def _download_image(self,url):
        try:
            headers = {'User-Agent':random.choice(self.ualist)}
            imagedata = requests.get(url,headers=headers).content
            return imagedata
        except:
            print('网络连接异常，可能爬取太多，被服务器暂封ip,稍后再试')

    def _getDatafromMongo(self):
        try:
            all_data = self.mongo_collection.find()
            return all_data
        except:
            print('数据库连接失败')

    def write_img(self,url):
        imgdata = self._download_image(url)
        image_io = BytesIO()
        image_io.write(imgdata)
        image_io.seek(0)
        return image_io

    def main(self):
        all_data = self._getDatafromMongo()
        for data in all_data:
            document = Document()
            doctext = data['内容']
            title = data['标题']
            auther = data['作者']
            time = data['发布时间']
            savename = re.sub(r'[\/:*?"<>|]','-',title)
            head = document.add_heading(title,0)
            head.add_run(title).bold = True
            p = document.add_heading(auther+'\n'+time)
            p.add_run(auther).italic = True
            for text in doctext:
                if text.startswith('http'):
                    imio = self.write_img(text)
                    im = Image.open(imio)
                    document.add_picture(imio,width=Inches(4))
                else:
                    p = document.add_paragraph(text)
            document.save(savename+'.docx')
            print('%s下载成功' %savename)

if __name__ == '__main__':
    x = Writeto_docx()
    x.main()



